import streamlit as st
import os
import json
import uuid
import base64
import shutil
import string
import random
import re
import requests
from datetime import datetime
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import io
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
import pyrebase
from streamlit_autorefresh import st_autorefresh

# ══════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="AM Intelligence",
    page_icon="🔒",
    layout="wide",
)

# ══════════════════════════════════════════════════════════════════════
# ENX LLM GATEWAY CONFIG
# The gateway is used for both chat (LLM) and embeddings.
# Azure credentials below are optional and used ONLY as a vision fallback
# (image queries) if the gateway does not support multimodal input.
# ══════════════════════════════════════════════════════════════════════
GATEWAY_BASE_URL    = "https://alb.prod.shared.aifoundations.digital.cloud.int/gateway"
GATEWAY_COMPANY     = st.secrets["gateway"]["company"]
GATEWAY_PROJECT_ID  = st.secrets["gateway"]["project_id"]
GATEWAY_TOKEN       = st.secrets["gateway"]["project_token"]
GATEWAY_CHAT_MODEL  = st.secrets["gateway"].get("chat_model", "gpt-5.2")
GATEWAY_CHAT_PROVIDER = st.secrets["gateway"].get("chat_provider")

_gateway_available_models = st.secrets["gateway"].get("chat_available_models", [])
if isinstance(_gateway_available_models, str):
    _gateway_available_models = [
        model.strip()
        for model in _gateway_available_models.split(",")
        if model.strip()
    ]

_gateway_chat_fallbacks = st.secrets["gateway"].get("chat_fallback_models", [])
if isinstance(_gateway_chat_fallbacks, str):
    _gateway_chat_fallbacks = [
        model.strip()
        for model in _gateway_chat_fallbacks.split(",")
        if model.strip()
    ]

if not _gateway_available_models:
    _gateway_available_models = [
        GATEWAY_CHAT_MODEL,
        *_gateway_chat_fallbacks,
    ]

_GATEWAY_PROVIDER_CHOICES = {
    "auto": "Auto detect",
    "azureOpenAi": "Azure OpenAI",
    "bedrock": "Amazon Bedrock",
}
_BUILTIN_GATEWAY_CHAT_MODELS = [
    {"model": "gpt-4o",          "label": "gpt-4o"},
    {"model": "gpt-4.1",         "label": "gpt-4.1"},
    {"model": "gpt-4.1-mini",    "label": "gpt-4.1-mini"},
    {"model": "gpt-5",           "label": "gpt-5"},
    {"model": "gpt-5.2",         "label": "gpt-5.2"},
    {"model": "gpt-5.2-chat",    "label": "gpt-5.2-chat"},
    {"model": "gpt-5.2-codex",   "label": "gpt-5.2-codex"},
    {
        "model": "claude-opus-4.6",
        "label": "claude-opus-4.6",
        "provider": "bedrock",
    },
    {
        "model": "claude-sonnet-4.6",
        "label": "claude-sonnet-4.6",
        "provider": "bedrock",
    },
    {
        "model": "claude-haiku-4.5",
        "label": "claude-haiku-4.5",
        "provider": "bedrock",
    },
]


def _normalize_gateway_provider(provider: str | None) -> str | None:
    if provider is None:
        return None
    normalized = str(provider).strip()
    if not normalized or normalized.lower() == "auto":
        return None
    return normalized


def _gateway_model_display_name(model: str, provider: str | None = None) -> str:
    provider_labels = {
        **_GATEWAY_PROVIDER_CHOICES,
        "AZURE_OPEN_AI": "Azure OpenAI",
        "BEDROCK": "Amazon Bedrock",
        "azureOpenAi": "Azure OpenAI",
        "bedrock": "Amazon Bedrock",
        "Azure OpenAI": "Azure OpenAI",
        "Amazon Bedrock": "Amazon Bedrock",
    }
    provider_label = provider_labels.get(provider)
    if provider_label and provider != "auto":
        return f"{model} ({provider_label})"
    return model


def _coerce_gateway_chat_model_spec(
    value,
    default_provider: str | None = None,
) -> dict | None:
    if isinstance(value, dict):
        model = str(value.get("model") or value.get("name") or "").strip()
        provider = _normalize_gateway_provider(
            value.get("provider") if "provider" in value else default_provider
        )
        label = str(value.get("label") or model).strip()
        if not model:
            return None
        return {
            "model": model,
            "provider": provider,
            "label": label or _gateway_model_display_name(model, provider),
        }

    if not isinstance(value, str):
        return None

    raw = value.strip()
    if not raw:
        return None

    provider = default_provider
    model = raw
    prefix, sep, remainder = raw.partition(":")
    if sep and prefix in _GATEWAY_PROVIDER_CHOICES and prefix != "auto":
        provider = prefix
        model = remainder.strip()

    builtin = next(
        (
            spec
            for spec in _BUILTIN_GATEWAY_CHAT_MODELS
            if spec["model"].lower() == model.lower()
        ),
        None,
    )
    if builtin:
        return dict(builtin)

    provider = _normalize_gateway_provider(provider)
    return {
        "model": model,
        "provider": provider,
        "label": _gateway_model_display_name(model, provider),
    }


def _dedupe_gateway_chat_model_specs(specs: list[dict]) -> list[dict]:
    deduped = []
    seen = set()
    for spec in specs:
        if not spec or not spec.get("model"):
            continue
        key = (
            spec["model"].lower(),
            (spec.get("provider") or "").lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(spec)
    return deduped


def _build_gateway_chat_model_catalog() -> list[dict]:
    configured = []
    for value in _gateway_available_models:
        spec = _coerce_gateway_chat_model_spec(
            value,
            default_provider=GATEWAY_CHAT_PROVIDER,
        )
        if spec:
            configured.append(spec)

    return _dedupe_gateway_chat_model_specs(configured)


GATEWAY_CHAT_MODEL_SPECS = _build_gateway_chat_model_catalog()
GATEWAY_EMBED_MODEL = "text-embedding-ada-002"
GATEWAY_VERIFY_SSL  = (
    st.secrets["gateway"].get("ssl_cert_path")
    or st.secrets["gateway"].get("verify_ssl", True)
)

# How many past turns to send to the LLM (keeps prompt size bounded)
HISTORY_WINDOW = 10

# ── Optional Azure fallback (vision / images only) ───────────────────
# If MODELS_ENDPOINT and LLM_API_KEY are present in secrets, image queries
# are routed to Azure.  If not, images are still accepted but answered
# in text-only mode with a warning.
_AZURE_ENDPOINT = st.secrets.get("MODELS_ENDPOINT")
_AZURE_API_KEY  = st.secrets.get("LLM_API_KEY")
_AZURE_API_VER  = "2025-04-01-preview"
_AZURE_LLM_DEP  = st.secrets.get("AZURE_LLM_DEPLOYMENT",
                                   "Llama-4-Maverick-17B-128E-Instruct-FP8")
AZURE_VISION_OK = bool(_AZURE_ENDPOINT and _AZURE_API_KEY)

# ── Regex helpers for page / section routing ─────────────────────────
PAGE_QUERY_RE       = re.compile(
    r"\bpage(?:\s+number|\s+no\.?)?\s+(\d+)\b", re.IGNORECASE
)
SECTION_TITLE_RE    = re.compile(
    r'\bsection\s+["""\'"]([^"""\'\"]+)["""\'"]', re.IGNORECASE
)
SECTION_FOLLOWUP_RE = re.compile(r"\b(that|this|the)\s+section\b", re.IGNORECASE)
MODEL_IDENTITY_RE   = re.compile(
    r"(?:\b(?:which|what)\s+(?:llm|model)\s+(?:are you|is this|am i using|powers you)\b"
    r"|\bwho\s+(?:made|created)\s+you\b"
    r"|\bare\s+you\s+(?:chatgpt|gpt[-\s]?4|gpt[-\s]?5)\b)",
    re.IGNORECASE,
)

# ── File upload limits ───────────────────────────────────────────────
MAX_UPLOAD_MB    = 20
MAX_UPLOAD_BYTES = MAX_UPLOAD_MB * 1024 * 1024

# ══════════════════════════════════════════════════════════════════════
# FIREBASE CONFIG
# ══════════════════════════════════════════════════════════════════════
FIREBASE_CONFIG = {
    "apiKey":            st.secrets["firebase"]["apiKey"],
    "authDomain":        st.secrets["firebase"]["authDomain"],
    "databaseURL":       st.secrets["firebase"]["databaseURL"],
    "projectId":         st.secrets["firebase"]["projectId"],
    "storageBucket":     st.secrets["firebase"]["storageBucket"],
    "messagingSenderId": st.secrets["firebase"]["messagingSenderId"],
    "appId":             st.secrets["firebase"]["appId"],
}

# ══════════════════════════════════════════════════════════════════════
# AUTHORIZED USERS
# ══════════════════════════════════════════════════════════════════════
AUTHORIZED_USERS = {
    "authuser": "password321",
}

# ══════════════════════════════════════════════════════════════════════
# LOCAL STORAGE PATHS
# ══════════════════════════════════════════════════════════════════════
APP_DATA_DIR = "./app_data"
SESSIONS_DIR = os.path.join(APP_DATA_DIR, "sessions")
ROOMS_DIR    = os.path.join(APP_DATA_DIR, "rooms")
os.makedirs(SESSIONS_DIR, exist_ok=True)
os.makedirs(ROOMS_DIR,    exist_ok=True)

# ══════════════════════════════════════════════════════════════════════
# SESSION STATE BOOTSTRAP
# ══════════════════════════════════════════════════════════════════════
DEFAULTS = {
    "authenticated":          False,
    "username":               "",
    "mode":                   "private",
    "current_session_id":     None,
    "current_room_code":      None,
    "pdf_stores":             {},
    "image_files":            {},
    "selected_files":         [],
    "chat_history":           [],
    "room_thinking":          False,
    "room_selected_files":    [],
    "private_upload_sig":     None,
    "room_upload_sig":        None,
    "room_unavailable_files": [],
    "gateway_working_chat_model": None,
    "gateway_working_chat_provider": None,
    "gateway_last_response_model": None,
    "gateway_last_response_provider": None,
    "gateway_selected_chat_model": GATEWAY_CHAT_MODEL,
}
for key, val in DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ══════════════════════════════════════════════════════════════════════
# FIREBASE INIT
# The full firebase_app is kept (not just db) because Firebase Storage
# is used to sync room files across devices.
# ══════════════════════════════════════════════════════════════════════
@st.cache_resource
def init_firebase():
    return pyrebase.initialize_app(FIREBASE_CONFIG)

firebase_app = init_firebase()
db           = firebase_app.database()

# ══════════════════════════════════════════════════════════════════════
# GATEWAY — HELPERS
# ══════════════════════════════════════════════════════════════════════
def _gateway_headers() -> dict:
    """Standard headers required by every gateway request."""
    return {
        "Content-Type": "application/json",
        "Accept":       "application/json",
        "company":      GATEWAY_COMPANY,
        "projectId":    GATEWAY_PROJECT_ID,
        "projectToken": GATEWAY_TOKEN,
    }


def _gateway_error_message(resp) -> str:
    try:
        payload = resp.json()
    except ValueError:
        return resp.text.strip() or f"HTTP {resp.status_code}"

    if isinstance(payload, dict):
        message = payload.get("message")
        if message:
            return message
        error = payload.get("error")
        if isinstance(error, str) and error:
            return error

    return str(payload)


def _selected_gateway_chat_spec() -> dict | None:
    selection = st.session_state.get(
        "gateway_selected_chat_model",
        GATEWAY_CHAT_MODEL,
    )

    if selection:
        for spec in GATEWAY_CHAT_MODEL_SPECS:
            if spec["model"] == selection:
                return spec

    return None


def _gateway_chat_model_candidates() -> list[dict]:
    selected = _selected_gateway_chat_spec()
    if selected:
        return [selected]

    preferred_model = st.session_state.get("gateway_working_chat_model")
    preferred_provider = st.session_state.get("gateway_working_chat_provider")
    candidates = []
    if preferred_model:
        candidates.append({
            "model": preferred_model,
            "provider": preferred_provider,
            "label": _gateway_model_display_name(
                preferred_model,
                preferred_provider,
            ),
        })

    return _dedupe_gateway_chat_model_specs([
        *candidates,
        *GATEWAY_CHAT_MODEL_SPECS,
    ])


def _format_gateway_chat_spec(spec: dict) -> str:
    return _gateway_model_display_name(spec["model"], spec.get("provider"))


def _current_gateway_model_status() -> str:
    selected = _selected_gateway_chat_spec()
    if selected:
        return selected["label"]

    working_model = st.session_state.get("gateway_working_chat_model")
    working_provider = st.session_state.get("gateway_working_chat_provider")
    if working_model:
        return _gateway_model_display_name(working_model, working_provider)

    if GATEWAY_CHAT_MODEL_SPECS:
        return GATEWAY_CHAT_MODEL_SPECS[0]["label"]

    return GATEWAY_CHAT_MODEL


def _effective_user_question(prompt: str) -> str:
    question_marker = "Question:"
    answer_marker = "\n\nAnswer:"
    if question_marker in prompt:
        extracted = prompt.rsplit(question_marker, 1)[-1]
        if answer_marker in extracted:
            extracted = extracted.split(answer_marker, 1)[0]
        return extracted.strip()
    return prompt.strip()


def _is_model_identity_question(prompt: str) -> bool:
    return bool(MODEL_IDENTITY_RE.search(_effective_user_question(prompt)))


def _model_identity_response(prompt: str, images: list = None) -> str | None:
    if not _is_model_identity_question(prompt):
        return None

    selected = _selected_gateway_chat_spec()
    selected_label = selected["label"] if selected else _current_gateway_model_status()
    backend_model = st.session_state.get("gateway_last_response_model")
    backend_provider = st.session_state.get("gateway_last_response_provider")

    parts = [
        f"This AM Intelligence chat is currently configured to use {selected_label}."
    ]
    if backend_model:
        parts.append(
            "The last successful gateway response for this chat reported "
            f"{_gateway_model_display_name(backend_model, backend_provider)}."
        )
    else:
        parts.append(
            "The gateway has not yet reported a backend-resolved model name for this chat."
        )
    if images:
        parts.append(
            f"If you attach images, the app can use the Azure vision fallback deployment {_AZURE_LLM_DEP}."
        )
    parts.append(
        "If you ask the model directly which model it is, it can answer inaccurately because that reply is generated text, not app metadata."
    )
    return " ".join(parts)


def _is_model_not_allowed(resp, model: str) -> bool:
    message = _gateway_error_message(resp).lower()
    return (
        resp.status_code == 400
        and "not allowed" in message
        and model.lower() in message
    )


def _extract_gateway_embedding_outputs(payload: dict) -> list[list[float]]:
    if isinstance(payload, dict):
        results = payload.get("results")
        if isinstance(results, list):
            outputs = []
            for item in results:
                if isinstance(item, dict) and isinstance(item.get("output"), list):
                    outputs.append(item["output"])
            if outputs:
                return outputs

        result = payload.get("result")
        if isinstance(result, dict) and isinstance(result.get("output"), list):
            return [result["output"]]

        output = payload.get("output")
        if isinstance(output, list):
            return [output]

        raise RuntimeError(
            "Gateway embedding response did not include any vectors. "
            f"Top-level keys: {', '.join(payload.keys()) or 'none'}."
        )

    raise RuntimeError(
        "Gateway embedding response was not a JSON object. "
        f"Received: {type(payload).__name__}."
    )


def _extract_gateway_chat_text(payload: dict) -> str:
    generations = payload.get("generations") or []
    if generations:
        assistant_message = generations[0].get("assistant_message") or {}
        text = assistant_message.get("text")
        if isinstance(text, str):
            return text
    raise RuntimeError(
        "Gateway chat response did not include assistant text."
    )


def _extract_gateway_chat_metadata(payload: dict) -> dict:
    metadata = payload.get("chat_response_metadata") or {}
    return {
        "model": metadata.get("model"),
        "provider": metadata.get("provider"),
    }

# ══════════════════════════════════════════════════════════════════════
# GATEWAY — EMBEDDINGS
# LangChain-compatible embeddings backed by the ENX LLM Gateway.
# embed_documents() → called when indexing uploaded file chunks.
# embed_query()     → called when converting a user question to a vector.
# ══════════════════════════════════════════════════════════════════════
class GatewayEmbeddings(Embeddings):
    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        if not texts:
            return []
        payload = {
            "inputs":  texts,
            "options": {"model": GATEWAY_EMBED_MODEL},
        }
        resp = requests.post(
            f"{GATEWAY_BASE_URL}/embedding",
            headers=_gateway_headers(),
            json=payload,
            verify=GATEWAY_VERIFY_SSL,
            timeout=60,
        )
        resp.raise_for_status()
        return _extract_gateway_embedding_outputs(resp.json())

    def embed_query(self, text: str) -> list[float]:
        payload = {
            "inputs":  [text],
            "options": {"model": GATEWAY_EMBED_MODEL},
        }
        resp = requests.post(
            f"{GATEWAY_BASE_URL}/embedding",
            headers=_gateway_headers(),
            json=payload,
            verify=GATEWAY_VERIFY_SSL,
            timeout=60,
        )
        resp.raise_for_status()
        outputs = _extract_gateway_embedding_outputs(resp.json())
        return outputs[0]


@st.cache_resource
def load_embeddings():
    return GatewayEmbeddings()

embeddings = load_embeddings()

# ══════════════════════════════════════════════════════════════════════
# GATEWAY — CHAT (LLM HELPERS)
#
# Temperature guide:
# 0.0 ─── 0.2 ─── 0.5 ─── 0.7 ─── 1.0+
#  |        |       |       |        |
# Robotic Precise Balance Natural  Chaotic
#
# 0.4 → document Q&A  (precise, grounded)
# 0.85 → open chat    (natural, conversational)
# ══════════════════════════════════════════════════════════════════════
_SYSTEM_DOC = (
    "You are a helpful AI assistant. "
    "When document context is provided, base your answer solely on it and "
    "always cite which document your answer comes from. "
    "If the context does not contain enough information, say so clearly."
)
_SYSTEM_CHAT = (
    "You are a helpful, friendly AI assistant. "
    "Answer questions naturally and conversationally."
)


def _build_gateway_history(history: list[dict], system_prompt: str) -> list[dict]:
    """
    Convert standard {"role": "user"/"assistant", "content": "..."} history
    into the Gateway chat_history format.

    The system prompt is injected into the first user message so the model
    always has its instructions regardless of gateway system-message support.
    """
    gw = []
    first_user_seen = False
    for msg in history:
        role = "USER" if msg["role"] == "user" else "ASSISTANT"
        text = msg["content"]
        if role == "USER" and not first_user_seen:
            text = f"[System Instructions]: {system_prompt}\n\n{text}"
            first_user_seen = True
        entry = {
            "message_type": role,
            "text_content":  text,
            "media":         [],
            "metadata":      {},
        }
        if role == "ASSISTANT":
            entry["tool_calls"] = []
        gw.append(entry)
    return gw


def ask_llm(
    prompt:      str,
    has_context: bool = False,
    images:      list = None,
    history:     list = None,
) -> str:
    """
    Send a message to the ENX LLM Gateway and return the text response.

    Parameters
    ----------
    prompt      : The user's current message / RAG-enriched prompt.
    has_context : True when document excerpts are embedded in the prompt.
    images      : List of base64 image strings.
                  Routes to the Azure vision fallback when present.
    history     : List of {"role": ..., "content": ...} prior messages.
                  Last HISTORY_WINDOW items are included.
    """
    system_prompt = _SYSTEM_DOC  if has_context else _SYSTEM_CHAT
    temperature   = 0.4          if has_context else 0.85

    identity_response = _model_identity_response(prompt, images=images)
    if identity_response:
        return identity_response

    # ── Vision path: route to Azure fallback ─────────────────────────
    if images:
        if AZURE_VISION_OK:
            return _ask_azure_with_images(prompt, has_context, images, history)
        else:
            st.warning(
                "⚠️ Vision/image queries require Azure credentials which are "
                "not configured. Answering in text-only mode."
            )
            images = None   # fall through to gateway text path

    # ── Build windowed history ────────────────────────────────────────
    windowed   = (history or [])[-HISTORY_WINDOW:]
    gw_history = _build_gateway_history(windowed, system_prompt) if windowed else []

    # Prepend system only when there is no prior history
    # (otherwise the system was already injected into the first history turn)
    current_text = (
        prompt if gw_history
        else f"[System Instructions]: {system_prompt}\n\n{prompt}"
    )

    payload = {
        "streaming":    False,
        "message": {
            "message_type": "USER",
            "text_content":  current_text,
            "media":         [],
            "metadata":      {},
        },
        "chat_history": gw_history,
        "temperature":  temperature,
    }

    selected_spec = _selected_gateway_chat_spec()

    rejected_models = []
    for spec in _gateway_chat_model_candidates():
        request_body = {**payload, "model": spec["model"]}
        if spec.get("provider"):
            request_body["provider"] = spec["provider"]

        resp = requests.post(
            f"{GATEWAY_BASE_URL}/conversate",
            headers=_gateway_headers(),
            json=request_body,
            verify=GATEWAY_VERIFY_SSL,
            timeout=120,
        )
        if resp.ok:
            response_payload = resp.json()
            response_meta = _extract_gateway_chat_metadata(response_payload)
            if selected_spec is None:
                st.session_state["gateway_working_chat_model"] = spec["model"]
                st.session_state["gateway_working_chat_provider"] = spec.get("provider")
            if response_meta.get("model"):
                st.session_state["gateway_last_response_model"] = response_meta["model"]
                st.session_state["gateway_last_response_provider"] = response_meta.get("provider")
            return _extract_gateway_chat_text(response_payload)

        error_message = _gateway_error_message(resp)
        if selected_spec is None and _is_model_not_allowed(resp, spec["model"]):
            rejected_models.append(_format_gateway_chat_spec(spec))
            continue

        raise RuntimeError(
            "Gateway chat request failed "
            f"for {_format_gateway_chat_spec(spec)} "
            f"({resp.status_code}): {error_message}"
        )

    tried_models = ", ".join(
        _format_gateway_chat_spec(spec)
        for spec in _gateway_chat_model_candidates()
    )
    rejected_text = ", ".join(rejected_models) or "none"
    raise RuntimeError(
        "Gateway chat request failed because this project does not allow any of "
        f"the configured chat models. Tried: {tried_models}. Rejected: {rejected_text}."
    )


# ── Azure vision fallback (images only) ──────────────────────────────
@st.cache_resource
def _load_azure_client():
    from openai import AzureOpenAI
    return AzureOpenAI(
        azure_endpoint=_AZURE_ENDPOINT,
        api_key=_AZURE_API_KEY,
        api_version=_AZURE_API_VER,
    )


def _ask_azure_with_images(
    prompt:      str,
    has_context: bool,
    images:      list,
    history:     list = None,
) -> str:
    """Azure OpenAI vision fallback — used only for image-bearing queries."""
    system_prompt = _SYSTEM_DOC  if has_context else _SYSTEM_CHAT
    temperature   = 0.4          if has_context else 0.85
    client        = _load_azure_client()

    messages = [{"role": "system", "content": system_prompt}]
    for msg in (history or [])[-HISTORY_WINDOW:]:
        messages.append({"role": msg["role"], "content": msg["content"]})

    content = [{"type": "text", "text": prompt}]
    for img_b64 in images:
        content.append({
            "type":      "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
        })
    messages.append({"role": "user", "content": content})

    resp = client.chat.completions.create(
        model=_AZURE_LLM_DEP,
        messages=messages,
        max_tokens=1000,
        temperature=temperature,
    )
    st.session_state["gateway_last_response_model"] = _AZURE_LLM_DEP
    st.session_state["gateway_last_response_provider"] = "Azure OpenAI"
    return resp.choices[0].message.content


def _render_gateway_model_selector():
    st.subheader("🤖 Choose your Model")

    options = [spec["model"] for spec in GATEWAY_CHAT_MODEL_SPECS]
    labels = {spec["model"]: spec["label"] for spec in GATEWAY_CHAT_MODEL_SPECS}

    current = st.session_state.get("gateway_selected_chat_model", GATEWAY_CHAT_MODEL)
    if current not in options:
        current = GATEWAY_CHAT_MODEL if GATEWAY_CHAT_MODEL in options else options[0]
        st.session_state["gateway_selected_chat_model"] = current

    st.selectbox(
        "Chat model",
        options=options,
        format_func=lambda option: labels.get(option, option),
        key="gateway_selected_chat_model",
        help=(
            "Switch models in real time using the models allowed for this project."
        ),
    )

    st.caption(f"Current selection: {_current_gateway_model_status()}")


# ══════════════════════════════════════════════════════════════════════
# RAG CONTEXT HELPERS
# (page-aware retrieval, section look-up, context formatting)
# ══════════════════════════════════════════════════════════════════════

def _format_context(matches: list[Document]) -> str:
    blocks = []
    for match in matches:
        source      = match.metadata.get("source", "unknown")
        page_number = match.metadata.get("page_number")
        heading     = f"[From: {source}"
        if page_number is not None:
            heading += f", page {page_number}"
        heading += "]"
        blocks.append(f"{heading}\n{match.page_content}")
    return "\n\n".join(blocks)


def _extract_requested_page(user_question: str):
    match = PAGE_QUERY_RE.search(user_question)
    return int(match.group(1)) if match else None


def _normalize_lookup_text(text: str) -> str:
    normalized = text.lower()
    normalized = normalized.replace("–", "-").replace("—", "-")
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def _extract_explicit_section_title(text: str):
    match = SECTION_TITLE_RE.search(text)
    if match:
        return match.group(1).strip()
    return None


def _resolve_section_title(user_question: str, history: list = None):
    explicit = _extract_explicit_section_title(user_question)
    if explicit:
        return explicit
    if not SECTION_FOLLOWUP_RE.search(user_question):
        return None
    for turn in reversed(history or []):
        if turn.get("role") != "user":
            continue
        prior_explicit = _extract_explicit_section_title(turn.get("content", ""))
        if prior_explicit:
            return prior_explicit
    return None


def _store_has_page_metadata(store) -> bool:
    doc_map = getattr(getattr(store, "docstore", None), "_dict", {})
    return any("page_number" in doc.metadata for doc in doc_map.values())


def _get_sorted_store_docs(store, filename: str) -> list[Document]:
    doc_map = getattr(getattr(store, "docstore", None), "_dict", {})
    docs = []
    for doc in doc_map.values():
        if "source" not in doc.metadata:
            doc.metadata["source"] = filename
        docs.append(doc)
    return sorted(
        docs,
        key=lambda doc: (
            doc.metadata.get("page_number", 10**9),
            doc.metadata.get("chunk_index", 10**9),
        )
    )


def _get_page_chunks(store, filename: str, page_number: int) -> list[Document]:
    matches = []
    for doc in _get_sorted_store_docs(store, filename):
        if doc.metadata.get("page_number") != page_number:
            continue
        matches.append(doc)
    return sorted(matches, key=lambda doc: doc.metadata.get("chunk_index", 0))


def _find_section_context(store, filename: str, section_title: str) -> list[Document]:
    docs   = _get_sorted_store_docs(store, filename)
    needle = _normalize_lookup_text(section_title)

    anchor_index = None
    for index, doc in enumerate(docs):
        if needle in _normalize_lookup_text(doc.page_content):
            anchor_index = index
            break

    if anchor_index is None:
        return []

    anchor       = docs[anchor_index]
    anchor_page  = anchor.metadata.get("page_number")
    anchor_chunk = anchor.metadata.get("chunk_index", 0)

    if anchor_page is None:
        return docs[anchor_index:anchor_index + 6]

    section_docs = []
    for doc in docs:
        page_number = doc.metadata.get("page_number")
        if page_number is None or page_number < anchor_page or page_number > anchor_page + 2:
            continue
        if page_number == anchor_page and doc.metadata.get("chunk_index", 0) < anchor_chunk:
            continue
        section_docs.append(doc)
    return section_docs


def build_response(user_question: str, history: list = None) -> str:
    """
    Build context from selected files and return an LLM answer.

    Supports:
    - Page-specific retrieval  ("show me page 5")
    - Section retrieval        ("tell me about section 'Introduction'")
    - General similarity search
    - Image queries (routed to Azure vision fallback if credentials exist)

    In room mode uses room_selected_files; in private mode uses selected_files.
    Passes history so the LLM maintains conversational memory across turns.
    """
    if st.session_state["mode"] == "room":
        selected = st.session_state["room_selected_files"]
    else:
        selected = st.session_state["selected_files"]

    pdf_stores  = st.session_state["pdf_stores"]
    image_files = st.session_state["image_files"]

    sel_text = [f for f in selected if f in pdf_stores]
    sel_imgs = [f for f in selected if f in image_files]
    imgs_b64 = [image_files[f] for f in sel_imgs]

    if not sel_text and not sel_imgs:
        return ask_llm(user_question, has_context=False, history=history)

    requested_page         = _extract_requested_page(user_question)
    section_title          = _resolve_section_title(user_question, history=history)
    context                = ""
    direct_page_lookup     = False
    direct_section_lookup  = False

    if sel_text:
        # ── Page-specific lookup ──────────────────────────────────────
        if requested_page is not None:
            page_matches     = []
            page_aware_files = []
            for fname in sel_text:
                store = pdf_stores[fname]
                if not _store_has_page_metadata(store):
                    continue
                page_aware_files.append(fname)
                page_matches.extend(_get_page_chunks(store, fname, requested_page))

            if page_matches:
                context            = _format_context(page_matches)
                direct_page_lookup = True
            elif page_aware_files and not sel_imgs:
                filenames = ", ".join(page_aware_files)
                return (
                    f"I could not find indexed text for page {requested_page} in "
                    f"the selected document(s): {filenames}. Please check the "
                    f"page number and try again."
                )
            elif not page_aware_files and not sel_imgs:
                return (
                    "Page-specific lookup is not available for the selected "
                    "document(s) yet. Please re-upload the PDF so it can be "
                    "indexed with page numbers."
                )

        # ── Section-specific lookup ───────────────────────────────────
        if section_title and not context:
            section_matches = []
            for fname in sel_text:
                section_matches.extend(
                    _find_section_context(pdf_stores[fname], fname, section_title)
                )
            if section_matches:
                context               = _format_context(section_matches)
                direct_section_lookup = True

        # ── General similarity search ─────────────────────────────────
        if not context:
            all_matches     = []
            retrieval_query = user_question
            if section_title:
                retrieval_query = f"{section_title}\n{user_question}"
            for fname in sel_text:
                matches = pdf_stores[fname].similarity_search(
                    retrieval_query,
                    k=4 if section_title else 2,
                )
                for match in matches:
                    match.metadata["source"] = fname
                all_matches.extend(matches)
            context = _format_context(all_matches)

    # ── Build the final prompt ────────────────────────────────────────
    if context and sel_imgs:
        if direct_page_lookup:
            prompt = (
                f"Use the exact excerpts from page {requested_page} AND the "
                f"provided images to answer.\n\n"
                f"Document Context:\n{context}\n\n"
                f"Question: {user_question}\n\nAnswer:"
            )
        elif direct_section_lookup:
            prompt = (
                f'Use the exact excerpts from the section titled '
                f'"{section_title}" AND the provided images to answer.\n\n'
                f"Document Context:\n{context}\n\n"
                f"Question: {user_question}\n\nAnswer:"
            )
        else:
            prompt = (
                f"Use the document excerpts AND the provided images to answer.\n\n"
                f"Document Context:\n{context}\n\n"
                f"Question: {user_question}\n\nAnswer:"
            )
    elif context:
        if direct_page_lookup:
            prompt = (
                f"Use the following exact excerpts from page {requested_page} to "
                f"answer the question.\n\n"
                f"Context:\n{context}\n\n"
                f"Question: {user_question}\n\nAnswer:"
            )
        elif direct_section_lookup:
            prompt = (
                f'Use the following exact excerpts from the section titled '
                f'"{section_title}" to answer the question.\n\n'
                f"Context:\n{context}\n\n"
                f"Question: {user_question}\n\nAnswer:"
            )
        else:
            prompt = (
                f"Use the following document excerpts to answer the question.\n\n"
                f"Context:\n{context}\n\n"
                f"Question: {user_question}\n\nAnswer:"
            )
    else:
        prompt = (
            f"Use the provided images to answer the question.\n\n"
            f"Question: {user_question}\n\nAnswer:"
        )

    return ask_llm(prompt, has_context=True, images=imgs_b64 or None, history=history)


# ══════════════════════════════════════════════════════════════════════
# FILE PROCESSING
# ══════════════════════════════════════════════════════════════════════

def extract_pages_from_pdf(f) -> list[Document]:
    """Extract text page-by-page, preserving page_number metadata."""
    f.seek(0)
    reader    = PdfReader(f)
    page_docs = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            page_docs.append(Document(
                page_content=text,
                metadata={"page_number": page_number, "filetype": "pdf"},
            ))
    return page_docs


def extract_text_from_docx(f) -> str:
    f.seek(0)
    doc = DocxDocument(f)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def image_to_base64(f) -> str:
    f.seek(0)
    img = Image.open(f)
    if img.mode != "RGB":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def _chunk_documents(documents: list[Document], filename: str) -> list[Document]:
    """Split documents into chunks while preserving page metadata."""
    splitter = RecursiveCharacterTextSplitter(
        separators=["\\n\\n", "\\n", " "],
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len,
    )
    chunks = []
    for document in documents:
        doc_chunks = splitter.split_text(document.page_content)
        for chunk_index, chunk in enumerate(doc_chunks):
            metadata                = dict(document.metadata)
            metadata["source"]      = filename
            metadata["chunk_index"] = chunk_index
            chunks.append(Document(page_content=chunk, metadata=metadata))
    return chunks


def chunk_and_embed(documents: list[Document], filename: str, index_dir: str):
    """Embed chunked documents and save a FAISS index locally."""
    chunks = _chunk_documents(documents, filename)
    if not chunks:
        return None
    store     = FAISS.from_documents(chunks, embeddings)
    save_path = os.path.join(index_dir, filename)
    os.makedirs(save_path, exist_ok=True)
    store.save_local(save_path)
    return store


def load_faiss_index(filename: str, index_dir: str):
    path = os.path.join(index_dir, filename)
    if os.path.exists(path):
        return FAISS.load_local(
            path, embeddings, allow_dangerous_deserialization=True
        )
    return None


def process_uploaded_file(
    uploaded_file,
    index_dir: str,
    mode: str = "private",
    session_id: str = None,
    room_code: str = None,
):
    """
    Process an uploaded file:
      - Extract text (PDF/DOCX) or base64-encode (images)
      - Embed and save to FAISS where applicable
      - Register metadata in session / room
      - Sync room files to Firebase Storage

    Notes:
    - room_selected_files is only updated when mode == "room" so that a
      private-session upload never leaks into room state.
    - Files are deduplicated; re-uploading a PDF that lacked page metadata
      will force a re-index.
    """
    fname     = uploaded_file.name
    ext       = fname.rsplit(".", 1)[-1].lower()
    file_size = getattr(uploaded_file, "size", 0)

    if file_size and file_size > MAX_UPLOAD_BYTES:
        size_mb = file_size / (1024 * 1024)
        raise ValueError(
            f"File is {size_mb:.1f} MB. Please upload a file under "
            f"{MAX_UPLOAD_MB} MB."
        )

    already_loaded = (
        list(st.session_state["pdf_stores"].keys()) +
        list(st.session_state["image_files"].keys())
    )
    if fname in already_loaded:
        existing_store = st.session_state["pdf_stores"].get(fname)
        # Allow re-upload of PDFs that were indexed without page metadata
        if not (ext == "pdf" and existing_store and not _store_has_page_metadata(existing_store)):
            st.info(f"'{fname}' is already loaded!")
            return

    if ext == "pdf":
        documents = extract_pages_from_pdf(uploaded_file)
        if not documents:
            raise ValueError(
                "No readable text could be extracted from this PDF. "
                "If it is a scanned document, OCR would be needed first."
            )
        store = chunk_and_embed(documents, fname, index_dir)
        if store:
            st.session_state["pdf_stores"][fname] = store
            if mode == "private" and session_id:
                _save_file_meta(session_id, fname, "pdf")
            elif mode == "room" and room_code:
                _add_file_to_room_db(room_code, fname, "pdf")
                if not _sync_room_file_to_cloud(room_code, fname, "pdf"):
                    st.warning(
                        f"{fname} was indexed locally, but shared room sync to "
                        f"Firebase Storage did not complete."
                    )

    elif ext == "docx":
        text = extract_text_from_docx(uploaded_file)
        if not text.strip():
            raise ValueError("No readable text could be extracted from this DOCX.")
        documents = [Document(page_content=text, metadata={"filetype": "docx"})]
        store = chunk_and_embed(documents, fname, index_dir)
        if store:
            st.session_state["pdf_stores"][fname] = store
            if mode == "private" and session_id:
                _save_file_meta(session_id, fname, "docx")
            elif mode == "room" and room_code:
                _add_file_to_room_db(room_code, fname, "docx")
                if not _sync_room_file_to_cloud(room_code, fname, "docx"):
                    st.warning(
                        f"{fname} was indexed locally, but shared room sync to "
                        f"Firebase Storage did not complete."
                    )

    elif ext in ["jpg", "jpeg", "png", "webp"]:
        img_b64 = image_to_base64(uploaded_file)
        st.session_state["image_files"][fname] = img_b64
        img_dir = (
            os.path.join(SESSIONS_DIR, session_id, "images")
            if mode == "private"
            else os.path.join(ROOMS_DIR, room_code, "images")
        )
        os.makedirs(img_dir, exist_ok=True)
        with open(os.path.join(img_dir, fname), "wb") as f:
            f.write(base64.b64decode(img_b64))
        if mode == "private" and session_id:
            _save_file_meta(session_id, fname, "image")
        elif mode == "room" and room_code:
            _add_file_to_room_db(room_code, fname, "image")
            if not _sync_room_file_to_cloud(room_code, fname, "image"):
                st.warning(
                    f"{fname} was saved locally, but shared room sync to "
                    f"Firebase Storage did not complete."
                )

    all_files = (
        list(st.session_state["pdf_stores"].keys()) +
        list(st.session_state["image_files"].keys())
    )
    st.session_state["selected_files"] = all_files
    if mode == "room":
        st.session_state["room_selected_files"] = all_files


def make_upload_signature(uploaded_file, scope_id: str) -> str:
    """Unique string to detect whether this file+scope has already been processed."""
    return f"{scope_id}:{uploaded_file.name}:{getattr(uploaded_file, 'size', 0)}"


# ══════════════════════════════════════════════════════════════════════
# FIREBASE STORAGE HELPERS (room file cloud sync)
# ══════════════════════════════════════════════════════════════════════

def _room_file_ref(code: str, filename: str):
    safe_key = _safe_firebase_key(filename)
    return db.child("rooms").child(code).child("files").child(safe_key)


def _room_index_dir(code: str, filename: str) -> str:
    return os.path.join(ROOMS_DIR, code, "indexes", filename)


def _room_image_path(code: str, filename: str) -> str:
    return os.path.join(ROOMS_DIR, code, "images", filename)


def _room_index_storage_base(code: str, filename: str) -> str:
    return f"room_assets/{code}/indexes/{_safe_firebase_key(filename)}"


def _room_image_storage_path(code: str, filename: str) -> str:
    return f"room_assets/{code}/images/{_safe_firebase_key(filename)}"


def _update_room_file_db(code: str, filename: str, updates: dict):
    _room_file_ref(code, filename).update(updates)


def _firebase_storage():
    return firebase_app.storage()


def _upload_to_storage(local_path: str, storage_path: str):
    _firebase_storage().child(storage_path).put(local_path)


def _download_from_storage(storage_path: str, local_path: str) -> bool:
    os.makedirs(os.path.dirname(local_path), exist_ok=True)
    _firebase_storage().child(storage_path).download(storage_path, local_path)
    return os.path.exists(local_path)


def _room_index_files_exist(code: str, filename: str) -> bool:
    index_dir = _room_index_dir(code, filename)
    return (
        os.path.exists(os.path.join(index_dir, "index.faiss")) and
        os.path.exists(os.path.join(index_dir, "index.pkl"))
    )


def _sync_room_file_to_cloud(code: str, filename: str, filetype: str) -> bool:
    try:
        if filetype in ["pdf", "docx"]:
            if not _room_index_files_exist(code, filename):
                return False
            storage_base = _room_index_storage_base(code, filename)
            index_dir    = _room_index_dir(code, filename)
            _upload_to_storage(
                os.path.join(index_dir, "index.faiss"),
                f"{storage_base}/index.faiss",
            )
            _upload_to_storage(
                os.path.join(index_dir, "index.pkl"),
                f"{storage_base}/index.pkl",
            )
            _update_room_file_db(code, filename, {
                "cloud_synced":  True,
                "storage_kind":  "faiss_index",
                "storage_base":  storage_base,
            })
            return True

        if filetype == "image":
            image_path = _room_image_path(code, filename)
            if not os.path.exists(image_path):
                return False
            storage_path = _room_image_storage_path(code, filename)
            _upload_to_storage(image_path, storage_path)
            _update_room_file_db(code, filename, {
                "cloud_synced":  True,
                "storage_kind":  "image",
                "storage_path":  storage_path,
            })
            return True
    except Exception:
        return False
    return False


def _download_room_file_from_cloud(code: str, filename: str, meta: dict) -> bool:
    try:
        filetype = meta.get("type", "")
        if filetype in ["pdf", "docx"]:
            storage_base = meta.get("storage_base") or _room_index_storage_base(code, filename)
            index_dir    = _room_index_dir(code, filename)
            os.makedirs(index_dir, exist_ok=True)
            faiss_ok = _download_from_storage(
                f"{storage_base}/index.faiss",
                os.path.join(index_dir, "index.faiss"),
            )
            pkl_ok = _download_from_storage(
                f"{storage_base}/index.pkl",
                os.path.join(index_dir, "index.pkl"),
            )
            return faiss_ok and pkl_ok

        if filetype == "image":
            storage_path = meta.get("storage_path") or _room_image_storage_path(code, filename)
            return _download_from_storage(storage_path, _room_image_path(code, filename))
    except Exception:
        return False
    return False


# ══════════════════════════════════════════════════════════════════════
# SESSION MANAGEMENT
# ══════════════════════════════════════════════════════════════════════

def get_all_sessions() -> list:
    sessions = []
    for sid in os.listdir(SESSIONS_DIR):
        mp = os.path.join(SESSIONS_DIR, sid, "metadata.json")
        if os.path.exists(mp):
            with open(mp) as f:
                sessions.append(json.load(f))
    return sorted(sessions, key=lambda x: x.get("updated_at", ""), reverse=True)


def create_new_session(name: str = None) -> str:
    sid         = str(uuid.uuid4())[:8]
    session_dir = os.path.join(SESSIONS_DIR, sid)
    os.makedirs(os.path.join(session_dir, "indexes"), exist_ok=True)
    os.makedirs(os.path.join(session_dir, "images"),  exist_ok=True)
    now  = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = {
        "id":         sid,
        "name":       name or f"New Chat · {now}",
        "created_at": now,
        "updated_at": now,
        "files":      [],
    }
    with open(os.path.join(session_dir, "metadata.json"), "w") as f:
        json.dump(meta, f, indent=2)
    with open(os.path.join(session_dir, "chat_history.json"), "w") as f:
        json.dump([], f)
    return sid


def load_session(sid: str):
    session_dir = os.path.join(SESSIONS_DIR, sid)
    hist_path   = os.path.join(session_dir, "chat_history.json")
    st.session_state["chat_history"] = (
        json.load(open(hist_path)) if os.path.exists(hist_path) else []
    )
    st.session_state["pdf_stores"]  = {}
    st.session_state["image_files"] = {}
    meta_path = os.path.join(session_dir, "metadata.json")
    if os.path.exists(meta_path):
        meta      = json.load(open(meta_path))
        index_dir = os.path.join(session_dir, "indexes")
        for fi in meta.get("files", []):
            fname, ftype = fi["name"], fi["type"]
            if ftype in ["pdf", "docx"]:
                store = load_faiss_index(fname, index_dir)
                if store:
                    st.session_state["pdf_stores"][fname] = store
            elif ftype == "image":
                img_path = os.path.join(session_dir, "images", fname)
                if os.path.exists(img_path):
                    with open(img_path, "rb") as f:
                        st.session_state["image_files"][fname] = (
                            base64.b64encode(f.read()).decode()
                        )
    st.session_state["selected_files"] = (
        list(st.session_state["pdf_stores"].keys()) +
        list(st.session_state["image_files"].keys())
    )
    st.session_state["current_session_id"] = sid


def save_chat_history(sid: str):
    session_dir = os.path.join(SESSIONS_DIR, sid)
    with open(os.path.join(session_dir, "chat_history.json"), "w") as f:
        json.dump(st.session_state["chat_history"], f, indent=2)
    meta_path = os.path.join(session_dir, "metadata.json")
    if os.path.exists(meta_path):
        meta = json.load(open(meta_path))
        meta["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        if meta["name"].startswith("New Chat") and st.session_state["chat_history"]:
            first        = st.session_state["chat_history"][0]["content"]
            meta["name"] = first[:45] + ("…" if len(first) > 45 else "")
        with open(meta_path, "w") as f:
            json.dump(meta, f, indent=2)


def _save_file_meta(sid: str, filename: str, filetype: str):
    meta_path = os.path.join(SESSIONS_DIR, sid, "metadata.json")
    meta      = json.load(open(meta_path))
    if not any(fi["name"] == filename for fi in meta["files"]):
        meta["files"].append({"name": filename, "type": filetype})
        with open(meta_path, "w") as f:
            json.dump(meta, f, indent=2)


def rename_session(sid: str, new_name: str):
    meta_path    = os.path.join(SESSIONS_DIR, sid, "metadata.json")
    meta         = json.load(open(meta_path))
    meta["name"] = new_name
    with open(meta_path, "w") as f:
        json.dump(meta, f, indent=2)


def delete_session(sid: str):
    path = os.path.join(SESSIONS_DIR, sid)
    if os.path.exists(path):
        shutil.rmtree(path)


# ══════════════════════════════════════════════════════════════════════
# ROOM MANAGEMENT
# ══════════════════════════════════════════════════════════════════════

def _generate_room_code() -> str:
    return "".join(random.choices(string.ascii_uppercase + string.digits, k=6))


def _safe_firebase_key(filename: str) -> str:
    """Firebase keys cannot contain dots, slashes, or spaces."""
    return filename.replace(".", "_").replace("/", "_").replace(" ", "_")


def create_room(room_name: str) -> str:
    code     = _generate_room_code()
    room_dir = os.path.join(ROOMS_DIR, code)
    os.makedirs(os.path.join(room_dir, "indexes"), exist_ok=True)
    os.makedirs(os.path.join(room_dir, "images"),  exist_ok=True)
    db.child("rooms").child(code).set({
        "name":       room_name,
        "created_by": st.session_state["username"],
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "active":     True,
        "files":      {},
    })
    return code


def room_exists(code: str) -> bool:
    try:
        val = db.child("rooms").child(code).get().val()
        return val is not None and val.get("active", False)
    except:
        return False


def get_room_messages(code: str) -> list:
    try:
        msgs = db.child("rooms").child(code).child("messages").get()
        if msgs.val():
            return sorted(
                msgs.val().items(),
                key=lambda x: x[1].get("timestamp", "")
            )
    except Exception as e:
        st.warning(f"{e}")
    return []


def send_room_message(
    code: str, user: str, content: str,
    is_bot: bool = False, is_system: bool = False,
):
    """
    Persist a message to the room on Firebase.
    is_bot    = True for LLM responses.
    is_system = True for notifications (file-selection changes etc.).
    """
    msg_id = str(uuid.uuid4())[:8]
    db.child("rooms").child(code).child("messages").child(msg_id).set({
        "user":      user,
        "content":   content,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "is_bot":    is_bot,
        "is_system": is_system,
    })


def get_room_files(code: str) -> list:
    """Return real filenames from the room's Firebase file registry."""
    registry = get_room_file_registry(code)
    return list(registry.keys())


def get_room_file_registry(code: str) -> dict:
    try:
        files = db.child("rooms").child(code).child("files").get()
        if files.val():
            return {
                v.get("filename", k): v
                for k, v in files.val().items()
                if isinstance(v, dict)
            }
        return {}
    except:
        return {}


def _room_history_for_llm(code: str) -> list:
    """
    Fetch room messages from Firebase, strip system notifications, and
    return the last HISTORY_WINDOW turns as {"role", "content"} dicts
    for the LLM history parameter.
    """
    try:
        raw = get_room_messages(code)
    except Exception:
        return []
    history = []
    for _, msg in raw:
        if msg.get("is_system"):
            continue
        role    = "assistant" if msg.get("is_bot") else "user"
        content = msg.get("content", "")
        if content:
            history.append({"role": role, "content": content})
    return history[-HISTORY_WINDOW:]


def _add_file_to_room_db(code: str, filename: str, filetype: str):
    safe_key = _safe_firebase_key(filename)
    db.child("rooms").child(code).child("files").child(safe_key).set({
        "type":         filetype,
        "filename":     filename,
        "uploaded_by":  st.session_state["username"],
        "uploaded_at":  datetime.now().strftime("%Y-%m-%d %H:%M"),
        "cloud_synced": False,
    })


def get_all_rooms() -> dict:
    try:
        rooms = db.child("rooms").get()
        if rooms.val():
            return {
                c: d for c, d in rooms.val().items()
                if d.get("active")
            }
        return {}
    except:
        return {}


def delete_room(code: str):
    db.child("rooms").child(code).update({"active": False})
    room_dir = os.path.join(ROOMS_DIR, code)
    if os.path.exists(room_dir):
        shutil.rmtree(room_dir)


def load_room_files_into_state(code: str):
    """
    Called on every auto-refresh so files uploaded by the authorized user
    become immediately available to all room members.

    - Backfills older files to Firebase Storage if cloud_synced is False.
    - Downloads files from Firebase Storage when the local index is missing.
    - Tracks unavailable files that could not be loaded.
    - Auto-selects all files for room querying by default.
    """
    room_files = get_room_file_registry(code)

    newly_loaded = []
    unavailable  = []
    for fname, meta in room_files.items():
        filetype = meta.get("type", "")

        # Backfill to Firebase Storage for older entries
        if not meta.get("cloud_synced"):
            _sync_room_file_to_cloud(code, fname, filetype)

        if (fname in st.session_state["pdf_stores"] or
                fname in st.session_state["image_files"]):
            continue

        if filetype in ["pdf", "docx"]:
            if not _room_index_files_exist(code, fname):
                if meta.get("cloud_synced") or meta.get("storage_base"):
                    _download_room_file_from_cloud(code, fname, meta)
            store = load_faiss_index(fname, os.path.join(ROOMS_DIR, code, "indexes"))
            if store:
                st.session_state["pdf_stores"][fname] = store
                newly_loaded.append(fname)
            else:
                unavailable.append(fname)

        elif filetype == "image":
            img_path = _room_image_path(code, fname)
            if not os.path.exists(img_path):
                if meta.get("cloud_synced") or meta.get("storage_path"):
                    _download_room_file_from_cloud(code, fname, meta)
            if os.path.exists(img_path):
                with open(img_path, "rb") as f:
                    st.session_state["image_files"][fname] = (
                        base64.b64encode(f.read()).decode()
                    )
                newly_loaded.append(fname)
            else:
                unavailable.append(fname)

    all_files = (
        list(st.session_state["pdf_stores"].keys()) +
        list(st.session_state["image_files"].keys())
    )
    st.session_state["selected_files"]         = all_files
    st.session_state["room_unavailable_files"] = unavailable

    if not st.session_state["room_selected_files"] and all_files:
        st.session_state["room_selected_files"] = all_files
    else:
        for f in newly_loaded:
            if f not in st.session_state["room_selected_files"]:
                st.session_state["room_selected_files"].append(f)


# ══════════════════════════════════════════════════════════════════════
# ██████████████████████████  S I D E B A R  ██████████████████████████
# ══════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.title("🔒 AM Intelligence")
    _render_gateway_model_selector()
    st.divider()

    # ── UNAUTHENTICATED ──────────────────────────────────────────────
    if not st.session_state["authenticated"]:

        if not st.session_state.get("current_room_code"):
            st.caption("💡 Join a group room with a code.")
            guest_code = st.text_input("Room Code", max_chars=6).upper().strip()
            guest_name = st.text_input("Your display name")
            if st.button("Join Room as Guest", use_container_width=True):
                if not guest_name:
                    st.error("Please enter a display name.")
                elif not room_exists(guest_code):
                    st.error("Room not found.")
                else:
                    st.session_state["mode"]                   = "room"
                    st.session_state["current_room_code"]      = guest_code
                    st.session_state["username"]               = guest_name
                    st.session_state["pdf_stores"]             = {}
                    st.session_state["image_files"]            = {}
                    st.session_state["selected_files"]         = []
                    st.session_state["room_selected_files"]    = []
                    st.session_state["room_unavailable_files"] = []
                    st.rerun()
        else:
            st.info(f"👤 **{st.session_state['username']}** (Guest)")
            st.caption(f"Room: **{st.session_state['current_room_code']}**")
            if st.button("Leave Room", use_container_width=True):
                st.session_state["mode"]                   = "private"
                st.session_state["current_room_code"]      = None
                st.session_state["username"]               = ""
                st.session_state["pdf_stores"]             = {}
                st.session_state["image_files"]            = {}
                st.session_state["selected_files"]         = []
                st.session_state["room_selected_files"]    = []
                st.session_state["room_unavailable_files"] = []
                st.rerun()

        st.divider()
        with st.expander("🔐 Login to upload files"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            if st.button("Login", use_container_width=True):
                if (username in AUTHORIZED_USERS and
                        AUTHORIZED_USERS[username] == password):
                    st.session_state["authenticated"] = True
                    st.session_state["username"]      = username
                    st.rerun()
                else:
                    st.error("Invalid credentials.")

    # ── AUTHENTICATED ────────────────────────────────────────────────
    else:
        st.success(f"👤 **{st.session_state['username']}**")

        mode_choice = st.radio(
            "App mode",
            ["💬 Private Chat", "👥 Group Room"],
            label_visibility="collapsed",
        )
        st.session_state["mode"] = (
            "private" if mode_choice == "💬 Private Chat" else "room"
        )

        if st.button("Logout", use_container_width=True):
            for k, v in DEFAULTS.items():
                st.session_state[k] = v
            st.rerun()

        st.divider()

        # ── PRIVATE CHAT SIDEBAR ─────────────────────────────────────
        if st.session_state["mode"] == "private":
            st.subheader("💬 Sessions")

            if st.button("➕ New Chat", use_container_width=True):
                sid = create_new_session()
                load_session(sid)
                st.rerun()

            for s in get_all_sessions():
                is_active = s["id"] == st.session_state.get("current_session_id")
                col1, col2 = st.columns([5, 1])
                with col1:
                    label = ("▶ " if is_active else "") + s["name"][:32]
                    if st.button(label, key=f"s_{s['id']}", use_container_width=True):
                        load_session(s["id"])
                        st.rerun()
                with col2:
                    if st.button("🗑", key=f"d_{s['id']}"):
                        delete_session(s["id"])
                        if s["id"] == st.session_state.get("current_session_id"):
                            st.session_state["current_session_id"] = None
                            st.session_state["chat_history"]       = []
                            st.session_state["pdf_stores"]         = {}
                            st.session_state["image_files"]        = {}
                        st.rerun()

            if st.session_state["current_session_id"]:
                st.divider()

                with st.expander("✏️ Rename session"):
                    new_name = st.text_input("New name", key="rename_inp")
                    if st.button("Save") and new_name:
                        rename_session(
                            st.session_state["current_session_id"], new_name
                        )
                        st.rerun()

                st.subheader("📁 Upload Files")
                st.caption(
                    f"Reliable indexing works best with files up to "
                    f"{MAX_UPLOAD_MB} MB."
                )
                uploaded = st.file_uploader(
                    "PDFs, DOCs or Images",
                    type=["pdf", "docx", "jpg", "jpeg", "png", "webp"],
                    key=f"priv_uploader_{st.session_state['current_session_id']}",
                )
                if uploaded:
                    sid         = st.session_state["current_session_id"]
                    index_dir   = os.path.join(SESSIONS_DIR, sid, "indexes")
                    current_sig = make_upload_signature(uploaded, sid)
                    if st.session_state["private_upload_sig"] != current_sig:
                        with st.spinner(f"Processing {uploaded.name}…"):
                            try:
                                process_uploaded_file(
                                    uploaded, index_dir,
                                    mode="private", session_id=sid,
                                )
                            except Exception as e:
                                st.error(f"Could not process {uploaded.name}: {e}")
                            else:
                                st.session_state["private_upload_sig"] = current_sig
                                st.success(f"✅ {uploaded.name} ready!")

                all_files = (
                    list(st.session_state["pdf_stores"].keys()) +
                    list(st.session_state["image_files"].keys())
                )
                if all_files:
                    st.subheader("🔍 Search in")
                    srch = st.radio(
                        "Search scope",
                        ["All files", "Select files"],
                        label_visibility="collapsed",
                    )
                    if srch == "All files":
                        st.session_state["selected_files"] = all_files
                    else:
                        st.session_state["selected_files"] = st.multiselect(
                            "Choose files",
                            options=all_files,
                            default=st.session_state["selected_files"] or all_files,
                        )

        # ── GROUP ROOM SIDEBAR ───────────────────────────────────────
        else:
            st.subheader("👥 Group Rooms")

            with st.expander("➕ Create New Room"):
                r_name = st.text_input("Room name", key="cr_name")
                if st.button("Create", use_container_width=True) and r_name:
                    code = create_room(r_name)
                    st.session_state["current_room_code"]      = code
                    st.session_state["pdf_stores"]             = {}
                    st.session_state["image_files"]            = {}
                    st.session_state["selected_files"]         = []
                    st.session_state["room_selected_files"]    = []
                    st.session_state["room_unavailable_files"] = []
                    st.success(f"✅ Room created!  Code: **{code}**")
                    st.rerun()

            with st.expander("🚪 Join a Room"):
                j_code = (
                    st.text_input("Room code", max_chars=6, key="j_code")
                    .upper().strip()
                )
                if st.button("Join", use_container_width=True, key="j_btn"):
                    if room_exists(j_code):
                        st.session_state["current_room_code"]      = j_code
                        st.session_state["pdf_stores"]             = {}
                        st.session_state["image_files"]            = {}
                        st.session_state["selected_files"]         = []
                        st.session_state["room_selected_files"]    = []
                        st.session_state["room_unavailable_files"] = []
                        st.rerun()
                    else:
                        st.error("Room not found.")

            all_rooms = get_all_rooms()
            my_rooms  = {
                c: d for c, d in all_rooms.items()
                if d.get("created_by") == st.session_state["username"]
            }
            if my_rooms:
                st.divider()
                st.caption("Your rooms:")
                for code, data in my_rooms.items():
                    is_active = code == st.session_state.get("current_room_code")
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        lbl = ("▶ " if is_active else "") + f"{data['name']} ({code})"
                        if st.button(lbl, key=f"r_{code}", use_container_width=True):
                            st.session_state["current_room_code"]      = code
                            st.session_state["pdf_stores"]             = {}
                            st.session_state["image_files"]            = {}
                            st.session_state["selected_files"]         = []
                            st.session_state["room_selected_files"]    = []
                            st.session_state["room_unavailable_files"] = []
                            st.rerun()
                    with col2:
                        if st.button("🗑", key=f"dr_{code}"):
                            delete_room(code)
                            if st.session_state.get("current_room_code") == code:
                                st.session_state["current_room_code"] = None
                            st.rerun()

            # ── File upload + file selector for active room ──────────
            if st.session_state.get("current_room_code"):
                st.divider()
                st.subheader("📁 Upload to Room")
                st.caption(
                    f"Reliable indexing works best with files up to "
                    f"{MAX_UPLOAD_MB} MB."
                )
                room_file = st.file_uploader(
                    "PDF, DOCX, or Image",
                    type=["pdf", "docx", "jpg", "jpeg", "png", "webp"],
                    key=f"room_uploader_{st.session_state['current_room_code']}",
                )
                if room_file:
                    code        = st.session_state["current_room_code"]
                    index_dir   = os.path.join(ROOMS_DIR, code, "indexes")
                    current_sig = make_upload_signature(room_file, code)
                    if st.session_state["room_upload_sig"] != current_sig:
                        with st.spinner(f"Processing {room_file.name}…"):
                            try:
                                process_uploaded_file(
                                    room_file, index_dir,
                                    mode="room", room_code=code,
                                )
                            except Exception as e:
                                st.error(f"Could not process {room_file.name}: {e}")
                            else:
                                st.session_state["room_upload_sig"] = current_sig
                                st.success(f"✅ {room_file.name} added to room!")

                # ── File selector (auth user only, visible in room) ──
                all_room_files = (
                    list(st.session_state["pdf_stores"].keys()) +
                    list(st.session_state["image_files"].keys())
                )
                if all_room_files:
                    st.subheader("🔍 Query scope")
                    room_srch = st.radio(
                        "Room search scope",
                        ["All files", "Select files"],
                        label_visibility="collapsed",
                        key="room_srch_radio",
                    )

                    prev_selection = list(st.session_state["room_selected_files"])

                    if room_srch == "All files":
                        st.session_state["room_selected_files"] = all_room_files
                    else:
                        st.session_state["room_selected_files"] = st.multiselect(
                            "Choose files",
                            options=all_room_files,
                            default=st.session_state["room_selected_files"] or all_room_files,
                            key="room_file_multiselect",
                        )

                    new_selection = st.session_state["room_selected_files"]
                    if set(new_selection) != set(prev_selection) and new_selection:
                        code  = st.session_state["current_room_code"]
                        notif = (
                            f"📌 **{st.session_state['username']}** changed the active "
                            f"documents to: **{', '.join(new_selection)}**"
                        )
                        send_room_message(code, "System", notif, is_system=True)


# ══════════════════════════════════════════════════════════════════════
# ████████████████████████  M A I N   A R E A  ████████████████████████
# ══════════════════════════════════════════════════════════════════════

mode = st.session_state["mode"]

# ── PRIVATE CHAT ─────────────────────────────────────────────────────
if mode == "private":

    st.header("🔒 AM Intelligence")
    st.caption(f"🤖 Model: {_current_gateway_model_status()}")

    if not st.session_state["authenticated"]:
        st.caption("💡 Log in from the sidebar to upload files and save sessions.")
        st.divider()

        for msg in st.session_state["chat_history"]:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        user_q = st.chat_input("Ask anything…")
        if user_q:
            with st.chat_message("user"):
                st.write(user_q)
            with st.spinner("Thinking…"):
                try:
                    response = ask_llm(
                        user_q, has_context=False,
                        history=st.session_state["chat_history"],
                    )
                except Exception as e:
                    response = f"⚠️ Error getting response: {str(e)}"
            with st.chat_message("assistant"):
                st.write(response)
            st.session_state["chat_history"].append({"role": "user",      "content": user_q})
            st.session_state["chat_history"].append({"role": "assistant", "content": response})

    elif not st.session_state["current_session_id"]:
        st.info("Click **➕ New Chat** in the sidebar to begin a saved session.")
        st.divider()

        for msg in st.session_state["chat_history"]:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        user_q = st.chat_input("Ask anything…")
        if user_q:
            with st.chat_message("user"):
                st.write(user_q)
            with st.spinner("Thinking…"):
                try:
                    response = ask_llm(
                        user_q, has_context=False,
                        history=st.session_state["chat_history"],
                    )
                except Exception as e:
                    response = f"⚠️ Error getting response: {str(e)}"
            with st.chat_message("assistant"):
                st.write(response)
            st.session_state["chat_history"].append({"role": "user",      "content": user_q})
            st.session_state["chat_history"].append({"role": "assistant", "content": response})

    else:
        sid       = st.session_state["current_session_id"]
        meta_path = os.path.join(SESSIONS_DIR, sid, "metadata.json")
        meta      = json.load(open(meta_path))

        st.header(f"💬 {meta['name']}")
        st.caption(f"🤖 Model: {_current_gateway_model_status()}")
        st.caption(f"🕐 Created: {meta['created_at']}  ·  Last updated: {meta['updated_at']}")

        if st.session_state["selected_files"]:
            st.caption(f"📂 Searching in: {', '.join(st.session_state['selected_files'])}")

        st.divider()

        for msg in st.session_state["chat_history"]:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        user_q = st.chat_input("Ask anything…")
        if user_q:
            with st.chat_message("user"):
                st.write(user_q)
            with st.spinner("Thinking…"):
                try:
                    response = build_response(
                        user_q, history=st.session_state["chat_history"]
                    )
                except Exception as e:
                    response = f"⚠️ Error getting response: {str(e)}"
            with st.chat_message("assistant"):
                st.write(response)
            st.session_state["chat_history"].append({"role": "user",      "content": user_q})
            st.session_state["chat_history"].append({"role": "assistant", "content": response})
            save_chat_history(sid)

# ── GROUP ROOM ────────────────────────────────────────────────────────
elif mode == "room":
    code = st.session_state.get("current_room_code")

    if not code:
        st.title("👥 Group Rooms")
        if st.session_state["authenticated"]:
            st.info("Create or join a room from the sidebar.")
        else:
            st.info("Enter a room code in the sidebar to join.")

    else:
        # Read chat input FIRST and set room_thinking=True immediately so that
        # st_autorefresh is never armed in the same render as an active LLM
        # call (which would silently kill it and produce no bot response).
        user_q = st.chat_input("Ask the room…")
        if user_q:
            st.session_state["room_thinking"] = True

        # Pause auto-refresh while LLM is thinking
        if not st.session_state.get("room_thinking", False):
            st_autorefresh(interval=15000, key="room_refresh")

        # Validate room
        try:
            room_data = db.child("rooms").child(code).get().val()
            if not room_data or not room_data.get("active"):
                st.error("This room no longer exists.")
                st.session_state["current_room_code"] = None
                st.stop()
        except Exception as e:
            st.error(f"Could not connect to room: {e}")
            st.stop()

        # Sync new files from disk into session_state
        load_room_files_into_state(code)

        # Header
        st.header(f"👥 {room_data['name']}")
        st.caption(
            f"Room Code: **{code}**  ·  "
            f"Created by: {room_data['created_by']}  ·  "
            f"{room_data['created_at']}"
        )
        st.caption(f"🤖 Model: {_current_gateway_model_status()}")

        room_files = get_room_files(code)
        if room_files:
            st.caption(f"📂 Files in room: {', '.join(room_files)}")

        unavailable = st.session_state.get("room_unavailable_files", [])
        if unavailable:
            st.warning(
                "These room files are registered in Firebase, but this device "
                f"could not load them yet: {', '.join(unavailable)}"
            )

        active = st.session_state.get("room_selected_files", [])
        if active and set(active) != set(room_files):
            st.caption(f"🔍 Currently querying: {', '.join(active)}")

        st.divider()

        # Render all messages from Firebase
        for _, msg in get_room_messages(code):
            if msg.get("is_system"):
                st.info(msg.get("content", ""))
            else:
                role = "assistant" if msg.get("is_bot") else "user"
                with st.chat_message(role):
                    if not msg.get("is_bot"):
                        st.caption(
                            f"**{msg.get('user', 'unknown')}**  ·  "
                            f"{msg.get('timestamp', '')}"
                        )
                    st.write(msg.get("content", ""))

        # Process chat input already captured above
        if user_q:
            send_room_message(
                code, st.session_state["username"], user_q, is_bot=False
            )

            with st.spinner("Thinking…"):
                try:
                    room_hist = _room_history_for_llm(code)
                    response  = build_response(user_q, history=room_hist)
                except Exception as e:
                    response = f"⚠️ Error getting response: {str(e)}"

            send_room_message(code, "AM Intelligence", response, is_bot=True)

            st.session_state["room_thinking"] = False
            st.rerun()
